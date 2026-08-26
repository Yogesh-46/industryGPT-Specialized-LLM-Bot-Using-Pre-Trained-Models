"""Build Woolf-format DataPilot AI research paper as .docx from markdown."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "research_paper.md"
OUT_PATH = ROOT / "DataPilot_AI_Research_paper.docx"


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc, text, *, style=None, size=12, bold=False, italic=False, center=False, space_after=8, space_before=0):
    p = doc.add_paragraph()
    if style:
        p.style = style
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def heading_level(line: str) -> int | None:
    if line.startswith("###### "):
        return 6
    if line.startswith("##### "):
        return 5
    if line.startswith("#### "):
        return 4
    if line.startswith("### "):
        return 3
    if line.startswith("## "):
        return 2
    if line.startswith("# "):
        return 1
    return None


def add_table(doc, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = ""
            p = table.rows[i].cells[j].paragraphs[0]
            run = p.add_run(cell.strip())
            set_run_font(run, size=10, bold=(i == 0))
    doc.add_paragraph()


def parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(set(c) <= set("-: ") and c for c in raw):
            rows.append(raw)
        i += 1
    return rows, i


def convert():
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    i = 0
    in_code = False
    code_buf: list[str] = []
    first_h1 = True

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                add_para(doc, "\n".join(code_buf), size=10, space_after=10)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").replace(" ", "")) <= set("-:"):
            rows, i = parse_md_table(lines, i)
            add_table(doc, rows)
            continue

        lvl = heading_level(line)
        if lvl is not None:
            title = line.lstrip("#").strip()
            if first_h1 and lvl == 1:
                add_para(doc, title, size=16, bold=True, center=True, space_after=12, space_before=6)
                first_h1 = False
            elif lvl == 2:
                add_para(doc, title, size=14, bold=True, space_before=16, space_after=8)
            elif lvl == 3:
                add_para(doc, title, size=13, bold=True, space_before=12, space_after=6)
            else:
                add_para(doc, title, size=12, bold=True, italic=True, space_before=10, space_after=4)
            i += 1
            continue

        if line.startswith("**Author:**") or line.startswith("Master's in CS"):
            add_para(doc, line.replace("**", ""), size=12, italic=True, center=True, space_after=4)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        # bullets
        if line.lstrip().startswith(("- ", "1. ", "2. ", "3. ", "4. ", "5. ", "6. ")):
            body = line.lstrip()
            if body[:2] in ("- ", "* "):
                body = body[2:]
            else:
                body = body.split(" ", 1)[-1]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(body.replace("**", ""))
            set_run_font(run, size=12)
            i += 1
            continue

        clean = line.replace("**", "").replace("`", "")
        add_para(doc, clean, size=12, space_after=8)
        i += 1

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    convert()
